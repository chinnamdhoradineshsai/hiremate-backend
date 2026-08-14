import io
import re
import uuid
from datetime import datetime
from typing import List, Optional
from fastapi import APIRouter, Depends, UploadFile, File, Form, HTTPException, status
from app.api.deps import get_current_user, UserProfileContext
from app.services.ai.ai_gateway import ai_gateway
from app.core.supabase import (
    get_supabase,
    SupabaseStorageService,
    is_supabase_configured,
)
from app.schemas.schemas import ATSResponse

from app.core.rate_limiter import rate_limit_expensive_ai

router = APIRouter(prefix="/resume", tags=["Resume & ATS"])

def normalize_text(text: str) -> str:
    if not text:
        return ""
    # Normalize carriage returns
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # Collapse 3 or more consecutive newlines into 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Collapse horizontal spaces/tabs per line
    lines = [re.sub(r'[ \t]+', ' ', line).strip() for line in text.split('\n')]
    return "\n".join(lines).strip()

def extract_pdf_text(file_bytes: bytes) -> tuple[str, int]:
    """
    Extracts text from all pages of a PDF file using pypdf.
    Returns (normalized_text, page_count).
    """
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        page_count = len(reader.pages)
        page_texts = []
        for page in reader.pages:
            txt = page.extract_text() or ""
            if txt.strip():
                page_texts.append(txt)
        raw_text = "\n\n".join(page_texts)
        return normalize_text(raw_text), page_count
    except Exception as e:
        print(f"[PDF Extract Error]: {e}")
        return "", 0

def extract_docx_text(file_bytes: bytes) -> tuple[str, int]:
    """
    Extracts text from paragraphs and tables of a DOCX file using python-docx.
    Returns (normalized_text, table_count).
    """
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        parts = []
        
        # Extract paragraph text
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        
        # Extract table text (rows & cells)
        table_count = len(doc.tables)
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_cells:
                    parts.append(" | ".join(row_cells))

        raw_text = "\n".join(parts)
        return normalize_text(raw_text), table_count
    except Exception as e:
        print(f"[DOCX Extract Error]: {e}")
        return "", 0

def validate_extracted_resume_text(text: str, min_length: int = 50) -> bool:
    """
    Validates that extracted text is non-empty, above minimum length, and contains meaningful words.
    """
    if not text or not isinstance(text, str):
        return False
    clean = text.strip()
    if len(clean) < min_length:
        return False
    words = re.findall(r'\w+', clean)
    if len(words) < 10:
        return False
    return True

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    ext = filename.lower().split('.')[-1]
    if ext == 'pdf':
        text, _ = extract_pdf_text(file_content)
        return text
    if ext == 'docx':
        text, _ = extract_docx_text(file_content)
        return text
    try:
        return normalize_text(file_content.decode('utf-8', errors='ignore'))
    except Exception:
        return ""

@router.post("/analyze", response_model=ATSResponse, dependencies=[Depends(rate_limit_expensive_ai)])
async def upload_and_analyze_resume(
    file: UploadFile = File(None),
    raw_resume_text: Optional[str] = Form(None),
    target_company: Optional[str] = Form(None),
    target_role: Optional[str] = Form(None),
    job_description: Optional[str] = Form(""),
    current_user: UserProfileContext = Depends(get_current_user)
):
    """
    Supabase Storage & Database Resume ATS Analysis Endpoint.
    Stores PDF/DOCX file in Supabase Storage bucket 'resumes' and metadata in Supabase tables.
    """
    # 0. Validate required parameters with clear, actionable error messages
    if not target_company or not target_company.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Missing required field: target_company. Please specify target company."
        )
    if not target_role or not target_role.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Missing required field: target_role. Please specify target role."
        )

    target_company = target_company.strip()
    target_role = target_role.strip()
    job_description = (job_description or "").strip()

    extracted_text = ""
    filename = "pasted_resume.txt"
    file_bytes = b""
    content_type = "text/plain"
    page_or_table_count = 0
    file_type_label = "plain_text"

    if file:
        filename = file.filename
        content_type = file.content_type or "application/pdf"
        file_bytes = await file.read()
        ext = filename.lower().split('.')[-1]

        if ext == 'pdf':
            file_type_label = "PDF"
            extracted_text, page_or_table_count = extract_pdf_text(file_bytes)
        elif ext == 'docx':
            file_type_label = "DOCX"
            extracted_text, page_or_table_count = extract_docx_text(file_bytes)
        elif ext in ['txt', 'md']:
            file_type_label = "TXT"
            try:
                extracted_text = normalize_text(file_bytes.decode('utf-8', errors='ignore'))
            except Exception:
                extracted_text = ""
        else:
            file_type_label = f"UNKNOWN ({ext})"
            try:
                extracted_text = normalize_text(file_bytes.decode('utf-8', errors='ignore'))
            except Exception:
                extracted_text = ""

    if not validate_extracted_resume_text(extracted_text) and raw_resume_text:
        file_type_label = "raw_resume_text"
        extracted_text = normalize_text(raw_resume_text)
        file_bytes = extracted_text.encode('utf-8')

    is_valid = validate_extracted_resume_text(extracted_text)

    # Safe diagnostic logging ONLY (NO resume contents or secrets!)
    print(f"[Resume Diagnostics] File Type: {file_type_label} | Extracted Char Count: {len(extracted_text)} | Page/Table Count: {page_or_table_count} | Extraction Success: {is_valid}")

    if not is_valid:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unable to extract readable text from this resume. Please upload a text-based PDF or DOCX."
        )

    resume_id = str(uuid.uuid4())
    supabase = get_supabase() if is_supabase_configured() else None

    # 1. Upload Resume file to Supabase Storage Bucket 'resumes' if configured
    storage_path = f"{current_user.id}/{resume_id}/{filename}"
    if file_bytes and supabase:
        try:
            await SupabaseStorageService.upload_resume_file(
                user_id=current_user.id,
                resume_id=resume_id,
                filename=filename,
                file_bytes=file_bytes,
                content_type=content_type
            )
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to upload resume to storage: {str(e)}"
            )

    # 2. Persist Resume metadata in Supabase 'resumes' table
    resume_row = {
        "id": resume_id,
        "user_id": current_user.id,
        "filename": filename,
        "storage_path": storage_path,
        "file_type": content_type,
        "extracted_text": extracted_text,
        "company": target_company,
        "role": target_role,
        "uploaded_at": datetime.utcnow().isoformat()
    }
    if supabase:
        try:
            supabase.table("resumes").insert(resume_row).execute()
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to persist resume metadata in database: {str(e)}"
            )

    # 3. Execute ATS Analysis via AI Gateway
    try:
        ats_result = await ai_gateway.analyze_resume_ats(
            resume_text=extracted_text,
            target_company=target_company,
            target_role=target_role,
            job_description=job_description
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"AI ATS Analysis failed: {str(e)}. Please retry or switch AI provider."
        )

    # 4. Store ATS Analysis in Supabase 'ats_analyses' table
    analysis_id = str(uuid.uuid4())
    ats_row = {
        "id": analysis_id,
        "resume_id": resume_id,
        "user_id": current_user.id,
        "overall_score": ats_result.get("overall_score", 75),
        "formatting_score": 80,
        "keyword_score": 80,
        "skills_score": 80,
        "experience_score": 80,
        "education_score": 80,
        "breakdown": ats_result.get("breakdown", {}),
        "missing_skills": ats_result.get("missing_skills", []),
        "missing_keywords": ats_result.get("missing_keywords", []),
        "writing_improvements": ats_result.get("writing_improvements", []),
        "free_resources": ats_result.get("free_resources", []),
        "company": target_company,
        "role": target_role,
        "job_description": job_description,
        "created_at": datetime.utcnow().isoformat()
    }
    if supabase:
        try:
            supabase.table("ats_analyses").insert(ats_row).execute()
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to persist ATS analysis in database: {str(e)}"
            )

    # 5. Populate User Learning Items in Supabase 'learning_items' table
    learning_rows = []
    for idx, skill in enumerate(ats_result.get("missing_skills", [])):
        s_name = skill.get("skill_name", "Skill")
        res_info = next((r for r in ats_result.get("free_resources", []) if r.get("skill_name") == s_name), None)
        res_url = res_info.get("resource_url") if res_info else "https://www.freecodecamp.org/news/learn-data-structures-and-algorithms/"
        rtype = "company_role" if target_company and target_role else ("role" if target_role else "personalized")
        learning_rows.append({
            "id": str(uuid.uuid4()),
            "user_id": current_user.id,
            "skill_name": s_name,
            "category": "Technical",
            "priority": skill.get("importance", "High"),
            "resource_title": res_info.get("resource_title", f"Learn {s_name}") if res_info else f"Learn {s_name}",
            "resource_url": res_url,
            "difficulty": res_info.get("difficulty", "Intermediate") if res_info else "Intermediate",
            "source_name": res_info.get("source_name", "Verified Learning Guide") if res_info else "Verified Learning Guide",
            "status": "Pending",
            "roadmap_week": (idx % 4) + 1,
            "target_company": target_company,
            "target_role": target_role,
            "roadmap_type": rtype,
            "created_at": datetime.utcnow().isoformat()
        })

    if learning_rows and supabase:
        try:
            supabase.table("learning_items").insert(learning_rows).execute()
        except Exception as e:
            print(f"[Supabase Learning Items Warning]: {e}")

    return ATSResponse(
        id=analysis_id,
        resume_id=resume_id,
        overall_score=ats_row["overall_score"],
        breakdown=ats_row["breakdown"],
        missing_skills=ats_row["missing_skills"],
        missing_keywords=ats_row["missing_keywords"],
        writing_improvements=ats_row["writing_improvements"],
        free_resources=ats_row["free_resources"],
        created_at=datetime.fromisoformat(ats_row["created_at"])
    )

@router.get("/latest", response_model=ATSResponse)
async def get_latest_ats_analysis(
    current_user: UserProfileContext = Depends(get_current_user)
):
    if not is_supabase_configured():
        raise HTTPException(status_code=404, detail="No resume analysis found for current user.")

    supabase = get_supabase()
    res = supabase.table("ats_analyses").select("*").eq("user_id", current_user.id).order("created_at", desc=True).limit(1).execute()
    data = res.data if res and res.data else []

    if not data:
        raise HTTPException(status_code=404, detail="No resume analysis found for current user.")

    latest = data[0]
    return ATSResponse(
        id=latest["id"],
        resume_id=latest["resume_id"],
        overall_score=latest["overall_score"],
        breakdown=latest["breakdown"],
        missing_skills=latest["missing_skills"],
        missing_keywords=latest["missing_keywords"],
        writing_improvements=latest["writing_improvements"],
        free_resources=latest["free_resources"],
        created_at=datetime.fromisoformat(latest["created_at"]) if isinstance(latest["created_at"], str) else latest["created_at"]
    )

@router.get("/history", response_model=List[ATSResponse])
async def get_ats_history(
    current_user: UserProfileContext = Depends(get_current_user)
):
    if not is_supabase_configured():
        return []

    supabase = get_supabase()
    res = supabase.table("ats_analyses").select("*").eq("user_id", current_user.id).order("created_at", desc=True).execute()
    analyses = res.data if res and res.data else []

    return [
        ATSResponse(
            id=item["id"],
            resume_id=item["resume_id"],
            overall_score=item["overall_score"],
            breakdown=item["breakdown"],
            missing_skills=item["missing_skills"],
            missing_keywords=item["missing_keywords"],
            writing_improvements=item["writing_improvements"],
            free_resources=item["free_resources"],
            created_at=datetime.fromisoformat(item["created_at"]) if isinstance(item["created_at"], str) else item["created_at"]
        ) for item in analyses
    ]
